"""Professional document export for Contract Revenue Radar reports.

Supports Markdown (always) and DOCX (when python-docx is installed).
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from .core import AuditReport, render_markdown


def render_docx(report: AuditReport, title: str = "Contract Revenue Radar - Revenue Terms Audit Report") -> bytes:
    """Render the audit report as a .docx file bytes.

    Requires python-docx. Raises ImportError with helpful message if unavailable.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install with: pip install 'contract-revenue-radar[export]' or pip install python-docx"
        ) from exc

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata block
    meta = doc.add_paragraph()
    meta.add_run("Backend: ").bold = True
    meta.add_run(report.backend)
    meta.add_run("  |  ")
    meta.add_run("Sections analyzed: ").bold = True
    meta.add_run(str(report.searched_sections))
    meta.add_run("  |  ")
    meta.add_run("Revenue Risk Score: ").bold = True
    score_run = meta.add_run(f"{report.risk_score}/100")
    if report.risk_score >= 80:
        score_run.font.color.rgb = RGBColor(180, 40, 40)
    elif report.risk_score >= 50:
        score_run.font.color.rgb = RGBColor(180, 120, 30)

    doc.add_paragraph()

    # Priority Actions
    doc.add_heading("Priority Negotiation Actions", level=1)
    for action in report.top_actions:
        p = doc.add_paragraph(action, style="List Bullet")

    # Findings
    doc.add_heading("Detailed Findings", level=1)

    for finding in report.findings:
        # Finding header
        h = doc.add_heading(f"{finding.label} — Severity {finding.severity}/5 (score {finding.score})", level=2)

        # Meta details
        details = doc.add_paragraph()
        details.add_run("Source: ").bold = True
        details.add_run(f"{finding.source} / {finding.heading}")
        details.add_run("\n")
        details.add_run("Matched terms: ").bold = True
        details.add_run(", ".join(finding.matched_terms) if finding.matched_terms else "semantic match (vector)")

        # Why + Action
        why_p = doc.add_paragraph()
        why_p.add_run("Why it matters: ").bold = True
        why_p.add_run(finding.why)

        action_p = doc.add_paragraph()
        action_p.add_run("Recommended move: ").bold = True
        action_p.add_run(finding.action)

        # Excerpt as quote block
        excerpt_p = doc.add_paragraph()
        excerpt_p.add_run("Excerpt: ").italic = True
        excerpt_run = excerpt_p.add_run(f'"{finding.excerpt}"')
        excerpt_run.italic = True

        doc.add_paragraph()  # spacing

    # Footer note
    footer = doc.add_paragraph()
    footer.add_run("IMPORTANT: ").bold = True
    footer.add_run("This is a business-risk review for internal negotiation preparation only. "
                   "It is not legal advice. Final contract language must be reviewed and approved by qualified counsel.")
    footer.runs[-1].font.size = Pt(9)
    footer.runs[-1].font.italic = True

    # Write to bytes buffer
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_report(
    report: AuditReport,
    output_path: Path,
    format: str = "md",
    title: Optional[str] = None,
) -> Path:
    """Export report to file in requested format. Returns the written path.

    format: "md" or "docx" (case-insensitive).
    """
    fmt = format.lower().strip()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if fmt in ("md", "markdown"):
        content = render_markdown(report)
        if title:
            # simple prepend if desired
            content = f"# {title}\n\n" + content
        output_path = output_path.with_suffix(".md")
        output_path.write_text(content, encoding="utf-8")
        return output_path

    if fmt in ("docx", "word", "doc"):
        docx_bytes = render_docx(report, title=title or "Contract Revenue Radar - Revenue Terms Audit Report")
        output_path = output_path.with_suffix(".docx")
        output_path.write_bytes(docx_bytes)
        return output_path

    raise ValueError(f"Unsupported export format: {format}. Use 'md' or 'docx'.")


def has_docx_support() -> bool:
    """Quick check if python-docx is importable for CLI / UI hints."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False
